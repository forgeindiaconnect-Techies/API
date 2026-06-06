import os
import asyncio
import logging
from typing import Dict, Any, Optional

logger = logging.getLogger(__name__)


async def process_dataset(file_path: str, file_type: str) -> Dict[str, Any]:
    """Process uploaded dataset and extract metadata"""
    loop = asyncio.get_event_loop()
    return await loop.run_in_executor(None, _process_sync, file_path, file_type)


def _process_sync(file_path: str, file_type: str) -> Dict[str, Any]:
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"Dataset file not found at: {file_path}")

    if file_type in ("csv",):
        return _process_csv(file_path)
    elif file_type in ("xlsx", "xls"):
        return _process_excel(file_path)
    elif file_type == "pdf":
        return _process_pdf(file_path)
    elif file_type in ("txt", "md"):
        return _process_text(file_path)
    elif file_type == "docx":
        return _process_docx(file_path)
    elif file_type == "json":
        return _process_json(file_path)
    elif file_type in ("jpg", "jpeg", "png", "webp"):
        return _process_image(file_path)
    elif file_type in ("mp3", "wav", "m4a"):
        return _process_audio(file_path)
    
    raise ValueError(f"Unsupported file type: {file_type}")


def _process_csv(path: str) -> Dict[str, Any]:
    import pandas as pd
    
    encodings = ["utf-8", "latin-1", "utf-8-sig", "cp1252"]
    df = None
    for enc in encodings:
        try:
            df = pd.read_csv(path, encoding=enc)
            break
        except Exception:
            continue
            
    if df is None:
        df = pd.read_csv(path)
        
    return {
        "rows": len(df),
        "cols": len(df.columns),
        "columns": list(df.columns),
        "metadata": {
            "dtypes": {col: str(dtype) for col, dtype in df.dtypes.items()},
            "missing_total": int(df.isnull().sum().sum()),
            "memory_mb": round(df.memory_usage(deep=True).sum() / 1024 / 1024, 2),
        }
    }


def _process_excel(path: str) -> Dict[str, Any]:
    import pandas as pd
    df = pd.read_excel(path)
    return {
        "rows": len(df),
        "cols": len(df.columns),
        "columns": list(df.columns),
        "metadata": {
            "dtypes": {col: str(dtype) for col, dtype in df.dtypes.items()},
            "missing_total": int(df.isnull().sum().sum()),
            "sheets": 1,
        }
    }


def _process_pdf(path: str) -> Dict[str, Any]:
    try:
        import PyPDF2
        import re
        with open(path, "rb") as f:
            reader = PyPDF2.PdfReader(f)
            pages = len(reader.pages)
            text_sample = ""
            full_text = []
            for i in range(pages):
                page_text = reader.pages[i].extract_text()
                if page_text:
                    full_text.append(page_text)
                    if i == 0:
                        text_sample = page_text[:500]
            
            content = "\n".join(full_text)
            words = len(content.split())
            
            # Calculate average sentence length
            sentences = re.split(r'[.!?]+', content)
            sentence_lengths = [len(s.strip().split()) for s in sentences if s.strip().split()]
            avg_sentence_len = round(sum(sentence_lengths) / len(sentence_lengths), 1) if sentence_lengths else 0
            
            # Language detection
            lang_stopwords = {
                "en": {"the", "and", "with", "for", "this", "that", "from"},
                "es": {"el", "la", "los", "las", "del", "con", "para", "por", "que", "este"},
                "fr": {"le", "la", "les", "des", "avec", "pour", "dans", "cette"},
                "de": {"der", "die", "das", "und", "mit", "von", "den", "dem", "für"}
            }
            words_list = re.findall(r'\b\w{3,}\b', content.lower())
            lang_scores = {lang: 0 for lang in lang_stopwords}
            for w in words_list[:1000]:
                for lang, sw in lang_stopwords.items():
                    if w in sw:
                        lang_scores[lang] += 1
            detected_lang = max(lang_scores, key=lang_scores.get) if any(lang_scores.values()) else "en"
            lang_mapping = {"en": "English", "es": "Spanish", "fr": "French", "de": "German"}
            language_name = lang_mapping.get(detected_lang, "English")

        return {
            "rows": pages,
            "cols": None,
            "columns": [],
            "metadata": {
                "pages": pages,
                "words": words,
                "chars": len(content),
                "text_sample": text_sample,
                "avg_sentence_len": avg_sentence_len,
                "language": language_name
            },
        }
    except Exception as e:
        logger.error(f"PDF metadata extraction failed: {e}")
        return {"rows": None, "cols": None, "columns": [], "metadata": {"type": "pdf", "error": str(e)}}


def _process_text(path: str) -> Dict[str, Any]:
    with open(path, "r", encoding="utf-8", errors="ignore") as f:
        content = f.read()
    lines = content.splitlines()
    words = len(content.split())
    
    # Calculate average sentence length
    import re
    sentences = re.split(r'[.!?]+', content)
    sentence_lengths = [len(s.strip().split()) for s in sentences if s.strip().split()]
    avg_sentence_len = round(sum(sentence_lengths) / len(sentence_lengths), 1) if sentence_lengths else 0
    
    # Language detection
    lang_stopwords = {
        "en": {"the", "and", "with", "for", "this", "that", "from"},
        "es": {"el", "la", "los", "las", "del", "con", "para", "por", "que", "este"},
        "fr": {"le", "la", "les", "des", "avec", "pour", "dans", "cette"},
        "de": {"der", "die", "das", "und", "mit", "von", "den", "dem", "für"}
    }
    words_list = re.findall(r'\b\w{3,}\b', content.lower())
    lang_scores = {lang: 0 for lang in lang_stopwords}
    for w in words_list[:1000]:
        for lang, sw in lang_stopwords.items():
            if w in sw:
                lang_scores[lang] += 1
    detected_lang = max(lang_scores, key=lang_scores.get) if any(lang_scores.values()) else "en"
    lang_mapping = {"en": "English", "es": "Spanish", "fr": "French", "de": "German"}
    language_name = lang_mapping.get(detected_lang, "English")

    return {
        "rows": len(lines),
        "cols": None,
        "columns": [],
        "metadata": {
            "lines": len(lines),
            "words": words,
            "chars": len(content),
            "avg_sentence_len": avg_sentence_len,
            "language": language_name
        }
    }


def _process_docx(path: str) -> Dict[str, Any]:
    try:
        import docx
        import re
        doc = docx.Document(path)
        paragraphs = [p.text for p in doc.paragraphs if p.text]
        content = "\n".join(paragraphs)
        words = len(content.split())
        
        # Calculate average sentence length
        sentences = re.split(r'[.!?]+', content)
        sentence_lengths = [len(s.strip().split()) for s in sentences if s.strip().split()]
        avg_sentence_len = round(sum(sentence_lengths) / len(sentence_lengths), 1) if sentence_lengths else 0
        
        # Language detection
        lang_stopwords = {
            "en": {"the", "and", "with", "for", "this", "that", "from"},
            "es": {"el", "la", "los", "las", "del", "con", "para", "por", "que", "este"},
            "fr": {"le", "la", "les", "des", "avec", "pour", "dans", "cette"},
            "de": {"der", "die", "das", "und", "mit", "von", "den", "dem", "für"}
        }
        words_list = re.findall(r'\b\w{3,}\b', content.lower())
        lang_scores = {lang: 0 for lang in lang_stopwords}
        for w in words_list[:1000]:
            for lang, sw in lang_stopwords.items():
                if w in sw:
                    lang_scores[lang] += 1
        detected_lang = max(lang_scores, key=lang_scores.get) if any(lang_scores.values()) else "en"
        lang_mapping = {"en": "English", "es": "Spanish", "fr": "French", "de": "German"}
        language_name = lang_mapping.get(detected_lang, "English")

        return {
            "rows": len(paragraphs),
            "cols": None,
            "columns": [],
            "metadata": {
                "paragraphs": len(paragraphs),
                "words": words,
                "chars": len(content),
                "avg_sentence_len": avg_sentence_len,
                "language": language_name
            }
        }
    except Exception as e:
        logger.error(f"DOCX metadata extraction failed: {e}")
        return {"rows": None, "cols": None, "columns": [], "metadata": {"type": "docx", "error": str(e)}}


def _process_json(path: str) -> Dict[str, Any]:
    import json
    import pandas as pd
    with open(path, "r", encoding="utf-8", errors="ignore") as f:
        data = json.load(f)
    if isinstance(data, list):
        if data and isinstance(data[0], dict):
            df = pd.DataFrame(data)
            return {
                "rows": len(df),
                "cols": len(df.columns),
                "columns": list(df.columns),
                "metadata": {"type": "array_of_objects"},
            }
        else:
            return {
                "rows": len(data),
                "cols": 1,
                "columns": ["Value"],
                "metadata": {"type": "array"},
            }
    elif isinstance(data, dict):
        return {
            "rows": 1,
            "cols": len(data.keys()),
            "columns": list(data.keys()),
            "metadata": {"type": "object"},
        }
    return {"rows": None, "cols": None, "columns": [], "metadata": {}}


def _process_image(path: str) -> Dict[str, Any]:
    try:
        from PIL import Image
        with Image.open(path) as img:
            return {
                "rows": None,
                "cols": None,
                "columns": [],
                "metadata": {
                    "width": img.width,
                    "height": img.height,
                    "mode": img.mode,
                    "format": img.format,
                }
            }
    except Exception:
        return {"rows": None, "cols": None, "columns": [], "metadata": {}}


def _process_audio(path: str) -> Dict[str, Any]:
    return {
        "rows": None,
        "cols": None,
        "columns": [],
        "metadata": {"type": "audio", "path": os.path.basename(path)}
    }


async def run_eda(file_path: str, file_type: str) -> Dict[str, Any]:
    """Run Exploratory Data Analysis on tabular datasets"""
    loop = asyncio.get_event_loop()
    return await loop.run_in_executor(None, _eda_sync, file_path, file_type)


def _eda_sync(file_path: str, file_type: str) -> Dict[str, Any]:
    try:
        import pandas as pd
        import numpy as np
        import os

        if file_type in ("csv", "xlsx", "xls"):
            if file_type == "csv":
                df = pd.read_csv(file_path)
            else:
                df = pd.read_excel(file_path)

            numeric_cols = list(df.select_dtypes(include=[np.number]).columns)
            categorical_cols = list(df.select_dtypes(exclude=[np.number]).columns)
            missing_by_col = {col: int(df[col].isnull().sum()) for col in df.columns if df[col].isnull().any()}

            col_stats = {}
            for col in numeric_cols[:10]:
                col_stats[col] = {
                    "mean": round(float(df[col].mean()), 4),
                    "std": round(float(df[col].std()), 4),
                    "min": round(float(df[col].min()), 4),
                    "max": round(float(df[col].max()), 4),
                    "median": round(float(df[col].median()), 4),
                    "q25": round(float(df[col].quantile(0.25)), 4),
                    "q75": round(float(df[col].quantile(0.75)), 4),
                    "nulls": int(df[col].isnull().sum()),
                }

            for col in categorical_cols[:5]:
                col_stats[col] = {
                    "unique_values": int(df[col].nunique()),
                    "top_values": df[col].value_counts().head(5).to_dict(),
                    "nulls": int(df[col].isnull().sum()),
                }

            # Correlation matrix (numeric only, max 10 cols)
            corr_matrix = {}
            if len(numeric_cols) >= 2:
                corr = df[numeric_cols[:10]].corr().round(3)
                corr_matrix = corr.to_dict()

            return {
                "is_tabular": True,
                "rows": len(df),
                "cols": len(df.columns),
                "missing_values": int(df.isnull().sum().sum()),
                "missing_by_column": missing_by_col,
                "duplicates": int(df.duplicated().sum()),
                "numeric_columns": numeric_cols,
                "categorical_columns": categorical_cols,
                "column_stats": col_stats,
                "correlation_matrix": corr_matrix,
            }
            
        elif file_type in ("txt", "md", "pdf", "docx"):
            text = ""
            if file_type in ("txt", "md"):
                with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                    text = f.read()
            elif file_type == "pdf":
                import PyPDF2
                with open(file_path, "rb") as f:
                    reader = PyPDF2.PdfReader(f)
                    text = "\n".join([page.extract_text() or "" for page in reader.pages])
            elif file_type == "docx":
                import docx
                doc = docx.Document(file_path)
                text = "\n".join([p.text for p in doc.paragraphs if p.text])

            import re
            from collections import Counter

            # 1. Word frequency & keywords (excluding common stopwords)
            words = re.findall(r'\b\w{3,}\b', text.lower())
            stopwords = {"the", "and", "der", "die", "und", "ein", "eine", "das", "von", "mit", "den", "dem", "einem", "einer", "for", "with", "this", "that", "from", "you", "your", "they", "them", "are", "was", "were", "been", "have", "has", "had", "will", "would", "shall", "should", "can", "could", "may", "might", "must", "about", "above", "across", "after", "again", "against", "all", "almost", "along", "already", "also", "although", "always", "among", "amount", "another", "any", "anyhow", "anyone", "anything", "anyway", "anywhere"}
            filtered_words = [w for w in words if w not in stopwords]
            word_counts = Counter(filtered_words)
            top_words = dict(word_counts.most_common(15))
            top_keywords = [item[0] for item in word_counts.most_common(10)]

            # 2. Sentence length distribution
            sentences = re.split(r'[.!?]+', text)
            sentence_lengths = []
            for s in sentences:
                s_words = s.strip().split()
                if s_words:
                    sentence_lengths.append(len(s_words))

            bins = {
                "1-5 words": 0,
                "6-10 words": 0,
                "11-15 words": 0,
                "16-20 words": 0,
                "21-30 words": 0,
                "31+ words": 0
            }
            for length in sentence_lengths:
                if length <= 5:
                    bins["1-5 words"] += 1
                elif length <= 10:
                    bins["6-10 words"] += 1
                elif length <= 15:
                    bins["11-15 words"] += 1
                elif length <= 20:
                    bins["16-20 words"] += 1
                elif length <= 30:
                    bins["21-30 words"] += 1
                else:
                    bins["31+ words"] += 1

            total_sentences = len(sentence_lengths) or 1
            avg_sentence_len = round(sum(sentence_lengths) / total_sentences, 1)

            # 3. Simple language detection
            lang_stopwords = {
                "en": {"the", "and", "with", "for", "this", "that", "from"},
                "es": {"el", "la", "los", "las", "del", "con", "para", "por", "que", "este"},
                "fr": {"le", "la", "les", "des", "avec", "pour", "dans", "cette"},
                "de": {"der", "die", "das", "und", "mit", "von", "den", "dem", "für"}
            }
            lang_scores = {lang: 0 for lang in lang_stopwords}
            for w in words:
                for lang, sw in lang_stopwords.items():
                    if w in sw:
                        lang_scores[lang] += 1
            detected_lang = max(lang_scores, key=lang_scores.get) if any(lang_scores.values()) else "en"
            lang_mapping = {"en": "English", "es": "Spanish", "fr": "French", "de": "German"}
            language_name = lang_mapping.get(detected_lang, "English")

            return {
                "is_tabular": False,
                "rows": len(text.splitlines()),
                "cols": None,
                "missing_values": 0,
                "duplicates": 0,
                "metadata": {
                    "lines": len(text.splitlines()),
                    "words": len(text.split()),
                    "chars": len(text),
                    "avg_sentence_len": avg_sentence_len,
                    "language": language_name
                },
                "word_frequency": top_words,
                "top_keywords": top_keywords,
                "sentence_length_distribution": bins
            }

        elif file_type in ("jpg", "jpeg", "png", "webp"):
            from PIL import Image
            with Image.open(file_path) as img:
                w, h = img.size
                mode = img.mode
                fmt = img.format
            return {
                "is_tabular": False,
                "is_image": True,
                "rows": None,
                "cols": None,
                "missing_values": 0,
                "duplicates": 0,
                "metadata": {
                    "width": w,
                    "height": h,
                    "mode": mode,
                    "format": fmt
                }
            }

        else:
            return {"error": f"EDA not supported for file type: {file_type}"}

    except Exception as e:
        logger.error(f"EDA error: {e}")
        return {"error": str(e)}


def auto_preprocess(df, options: dict):
    """Auto preprocess dataframe"""
    import pandas as pd
    import numpy as np

    if options.get("remove_duplicates"):
        df = df.drop_duplicates()

    # Handle missing values
    strategy = options.get("handle_missing", "median")
    numeric_cols = df.select_dtypes(include=[np.number]).columns
    categorical_cols = df.select_dtypes(exclude=[np.number]).columns

    for col in numeric_cols:
        if df[col].isnull().any():
            if strategy == "median":
                df[col] = df[col].fillna(df[col].median())
            elif strategy == "mean":
                df[col] = df[col].fillna(df[col].mean())
            elif strategy == "zero":
                df[col] = df[col].fillna(0)
            elif strategy == "drop":
                df = df.dropna(subset=[col])

    for col in categorical_cols:
        if df[col].isnull().any():
            df[col] = df[col].fillna(df[col].mode()[0] if len(df[col].mode()) else "unknown")

    # Encode categorical
    if options.get("encode_categorical"):
        from sklearn.preprocessing import LabelEncoder
        le = LabelEncoder()
        for col in categorical_cols:
            if df[col].nunique() <= 20:
                df[col] = le.fit_transform(df[col].astype(str))

    # Normalize numeric
    if options.get("normalize"):
        from sklearn.preprocessing import StandardScaler
        scaler = StandardScaler()
        target = options.get("target_column")
        cols_to_scale = [c for c in numeric_cols if c != target]
        if cols_to_scale:
            df[cols_to_scale] = scaler.fit_transform(df[cols_to_scale])

    return df
